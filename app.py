import time
import os
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, make_response, send_file,Response
from flask_sqlalchemy import SQLAlchemy
from functools import wraps
import numpy as np
from video_capture import VideoCapture
# importing necessary libraries MAPPING.
import sqlite3
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
#for word files
import codecs
import io
from io import BytesIO
from docx import Document
#FACE RECOGNITION AND STUDENTS ATTENTION 
import sys
from werkzeug.routing import BaseConverter
import random
import face_recognition
import cv2
import torch
from torchvision import transforms
from PIL import Image
import cv2, queue, threading
from video_capture import VideoCapture
import torch
import datetime
import time
import pandas as pd
import numpy as np
from segmentation import get_yolov5
import json


conf_thresh = 0.5  # Confidence threshold for object detection
iou_thresh = 0.5  # IoU threshold for non-maximum suppression
img_size = 640  # Size of input image for detection
# Load YOLOv5 model
model = get_yolov5()
# Get names of classes to be detected
class_names = model.module.names if hasattr(model, 'module') else model.names

app = Flask(__name__)

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///db/database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = '8IR4M7-R3c74GjTHhKzWODaYVHuPGqn4w92DHLqeYJA'

db = SQLAlchemy(app)
# Import models here as to avoid circular import issue
from models import *

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/add_course', methods=['GET', 'POST'])
def add_course():
    if request.method == 'POST':
        # Retrieve form data
        course_name = request.form['course_name']
        teacher_email = request.form['teacher_email']
        students_email = request.form['students_email']
        program_id = request.form['program_id']

        # Check if course already exists in database
        course = Course.query.filter_by(course_name=course_name).first()
        if course is not None:
            # Check if course already assigned to this program
            if course.program_id == program_id:
                # Add new students to existing course
                course.students_email += students_email
                db.session.commit()
                flash('Course Entries Updated', 'success')
            else:
                flash('Course Already Exists in Another Program!', 'danger')
            return redirect(url_for('add_course'))

        # Create new course and save to database
        course = Course(
            course_name=course_name,
            teacher_email=teacher_email,
            students_email=students_email,
            program_id=program_id,
        )
        db.session.add(course)
        db.session.commit()
        flash('Course Successfully Added', 'success')
        return redirect(url_for('add_course'))

    # Retrieve all programs from database
    programs = Program.query.all()
    print(programs) 
    return render_template('add_course.html', programs=programs)



@app.route('/add_program', methods=['GET', 'POST'])
def add_program():
    if request.method == 'POST':
        # check if a program with the same name already exists
        data = Program.query.filter_by(name=request.form['name']).first()
        if data is not None:
            flash('Program with this name already exists!', 'danger')
            return redirect(url_for('add_program'))

        # read the Word file contents
        wordfile = request.files['wordfile']
        wordfile_content = wordfile.read()

        # create a new Program object and add to the database
        program = Program(
            name=request.form['name'],
            plo_clo=wordfile_content
        )
        db.session.add(program)
        db.session.commit()

        flash('Program added successfully', 'success')
        return redirect(url_for('add_program'))

    return render_template('add_program.html')

@app.route('/login_student', methods=['GET', 'POST'])
def login_student():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        data = Student.query.filter_by(email=email, password=password).first()

        if data is not None:
            student = Student.query.filter_by(
                email=email, password=password).one()
            session['std_logged_in'] = True
            uname = email.split('@')[0]
            session['uname'] = uname
            session['user_type'] = "student"
            session['email'] = email
            session['student_id'] = student.student_id
            session['name'] = student.name      
            flash('You are now logged in', 'success')
            return redirect(url_for('student'))
        else:
            error = 'Invalid login'
            return render_template('login_student.html', error=error)

    return render_template('login_student.html')


@app.route('/login_faculty', methods=['GET', 'POST'])
def login_faculty():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        data = Faculty.query.filter_by(email=email, password=password).first()
        if data is not None:
            faculty = Faculty.query.filter_by(
                email=email, password=password).one()
            if faculty.is_admin:
                session['is_admin'] = True
            else:
                session['is_admin'] = False

            session['fty_logged_in'] = True
            uname = email.split('@')[0]
            session['uname'] = uname
            session['email'] = email
            session['user_type'] ="faculty"
            session['name'] = faculty.name
            

            flash('You are now logged in', 'success')
            return redirect(url_for('faculty'))
        else:
            error = 'Invalid login'
            return render_template('login_faculty.html', error=error)

    return render_template('login_faculty.html')


def is_faculty_logged_in(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'fty_logged_in' in session:
            return f(*args, **kwargs)
        else:
            flash('Unauthorized, Please login!', 'danger')
            return redirect(url_for('login_faculty'))
    return wrap


def is_student_logged_in(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'std_logged_in' in session:
            return f(*args, **kwargs)
        else:
            flash('Unauthorized, Please login!', 'danger')
            return redirect(url_for('login_student'))
    return wrap

@app.route('/faculty')
@is_faculty_logged_in
def faculty():
    
    students_count = Student.query.with_entities(Student.student_id).count()
    faculty_count = Faculty.query.with_entities(Faculty.f_id).count()

    courses = Course.query.filter_by(teacher_email=session['email']).all()

    return render_template('dashboard_faculty.html', courses=courses, students_count=students_count, faculty_count=faculty_count)


@app.route('/faculty_registration', methods=['GET', 'POST'])
@is_faculty_logged_in
def register_faculty():
    if request.method == 'POST':
        data = Faculty.query.filter_by(email=request.form['email']).first()
        if 'isAdmin' in request.form:
            is_admin = True
        else:
            is_admin = False
        # if email does not already exist
        if data is None:
            faculty = Faculty(
                name=request.form['name'],
                email=request.form['email'],
                password=request.form['password'],
                is_admin=is_admin,
                registered_on=datetime.now()
            )
            db.session.add(faculty)
            db.session.commit()

            flash('Faculty registration successful', 'success')
            return render_template('register_faculty.html')
        else:
            flash('Faculty with this email already exists!', 'danger')
    return render_template('register_faculty.html')



@app.route('/student_registration', methods=['GET', 'POST'])
@is_faculty_logged_in
def register_student():
    if session['is_admin']:
        if request.method == 'POST':
            data = Student.query.filter_by(email=request.form['email']).first()
            if data is None:
                new_student = Student(
                    student_id=request.form['student_id'],
                    name=request.form['name'],
                    semester=request.form['semester'],
                    email=request.form['email'],
                    password=request.form['password'],
                    pic_path=f'static/images/users/{request.form["student_id"]}-{request.form["name"]}.jpg',
                    registered_on=datetime.now()
                )
                db.session.add(new_student)
                db.session.commit()

                if os.path.isfile('static/images/users/temp.jpg'):
                    os.rename('static/images/users/temp.jpg',
                            f'static/images/users/{request.form["student_id"]}-{request.form["name"]}.jpg')
                if 'img_captured' in session:
                    session.pop('img_captured')
                flash('Student registration successful', 'success')
                return render_template('register_student.html')
            else:
                flash('Student with this email already exists!', 'danger')

        if os.path.isfile('static/images/users/temp.jpg'):
            temp_pic = True
        else:
            temp_pic = False

        return render_template('register_student.html', temp_pic=temp_pic)
    else:
        flash('Admin Access Required!', 'danger')
        return render_template('dashboard_faculty.html')


@app.route("/capture_image")
@is_faculty_logged_in
def capture_image():
    session['dt'] = datetime.now()
    path = 'static/images/users'
    cap = cv2.VideoCapture(0)

    while True:
        # Capture frame-by-frame
        ret, frame = cap.read()

        # Display the resulting frame
        cv2.imshow('Press c to capture image', frame)
        if cv2.waitKey(1) & 0xFF == ord('c'):
            cv2.imwrite(os.path.join(path, 'temp.jpg'), frame)
            time.sleep(2)
            break

    # When everything done, release the capture
    cap.release()
    cv2.destroyAllWindows()

    session['img_captured'] = True

    return redirect(url_for('register_student'))


@app.route('/student')
@is_student_logged_in
def student():
    courses = Course.query.filter(Course.students_email.contains(session['email'])).all()
    courses_count = []
    # print(courses_count)

    
    for course in courses:
        courses_count.append(course.course_name)
    
    return render_template('dashboard_student.html', courses=courses, courses_count=courses_count)



@app.route('/download_plo_clo/<int:program_id>')
def download_plo_clo(program_id):
    program = Program.query.get_or_404(program_id)
    if program.plo_clo:
        # create a response with the file contents and appropriate headers
        document = Document(io.BytesIO(program.plo_clo))
        file_path = f'{program.name}_plo_clo.docx'
        document.save(file_path)
        return send_file(file_path, as_attachment=True)
    else:
        # handle case where plo_clo is not available
        flash('No learning outcomes available for this program.', 'warning')
        return redirect(url_for('dashboard_student.html'))
    
@app.route("/logout")
def logout():
    if 'std_logged_in' or 'fty_logged_in' in session:
        session.clear()
    return redirect(url_for('index'))

#######################################################################
#CLO PLO MAPPING
#######################################################################
def read_docx_table(document,table_num):
    table = document.tables[table_num-1]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    return df 
global_df = pd.DataFrame()

@app.route("/plo_clo_no", methods=['GET', 'POST'])
def program_mapping():
    # Connect to the SQLite database
    conn = sqlite3.connect('db/database.db')
    # Retrieve the blob data from the database
    cursor = conn.cursor()
    cursor.execute("SELECT program_id FROM Course WHERE course_name=?", (session['course_name'],))
    prog_id = cursor.fetchone()[0]
    cursor.execute("SELECT plo_clo FROM Program WHERE id=?", (prog_id,))
    blob_data = cursor.fetchone()[0]
    # Close the database connection
    conn.close()
    bytes_io = io.BytesIO(blob_data)
    document = Document(bytes_io) # extract tables from word document
    table_num=1
    df_plo = read_docx_table(document,table_num) # Creating a dataframe for PLOs and it will accept 'n' number of PLOs
    plo_copy = df_plo.copy()
    count_vectorizer = CountVectorizer(stop_words='english', min_df=0.005) # assigning count vectorizer breaking down a text into words
    # Data preprocessing for PLO dataframe
    plo_copy[1] = plo_copy[1].str.lower()
    corpus_plo = plo_copy[1].tolist() #it will seperate the rows into pairs 
    corpii_plo = count_vectorizer.fit_transform(corpus_plo)
    # extracting features names from PLO table
    feature_names_plo = count_vectorizer.get_feature_names_out()
    # Converting features to vector form and create a dataframe
    X1_plo = pd.DataFrame(corpii_plo.toarray(), columns=feature_names_plo)
    # Creating a dataframe for CLOs and it will accept 'n' number of CLOs
    table_num=2
    df_clo = read_docx_table(document,table_num)
    clo_copy = df_clo.copy()
    # Data preprocessing for CLO dataframe
    clo_copy[1] = clo_copy[1].str.lower()
    corpus_clo = clo_copy[1].tolist()
    corpii_clo = count_vectorizer.fit_transform(corpus_clo)
    # extracting features names from CLO table
    feature_names_clo = count_vectorizer.get_feature_names_out()
    # Converting features to vector form and create a dataframe
    X2_clo = pd.DataFrame(corpii_clo.toarray(), columns=feature_names_clo)
    # adding column index to the PLO table
    U1_plo = pd.concat([df_plo[0], X1_plo], axis=1).set_index(0)
    # adding column index to the CLO table
    U2_clo = pd.concat([df_clo[0], X2_clo], axis=1).set_index(0)
    append_words = list(map(str.lower,['Cite', 'Define', 'Describe', 'Draw', 'Enumerate', 'Identify' 'Index', 'Indicate', 'Label', 'List', 'Match', 'Meet', 'Name', 'Outline', 'Point', 'Quote', 'Read', 'Recall', 'Recite', 'Recognize', 'Record', 'Repeat', 'Reproduce','Review',
'Select', 'State', 'Study', 'Tabulate', 'Trace', 'Write', 'Add', 'Approximate', 'Articulate', 'Associate', 'Characterize', 'Clarify', 'Classify', 'Compare', 'Compute', 'Contrast', 'Convert', 'Defend', 'Detail', 'Differentiate',
'Discuss', 'Distinguish', 'Elaborate', 'Estimate', 'Example', 'Explain', 'Express', 'Extend', 'Extrapolate', 'Factor', 'Generalize', 'Give', 'Infer', 'Interact', 'Interpolate', 'Interpret', 'Observe', 'Paraphrase', 'Picture graphically',
'Predict', 'Rewrite', 'Subtract', 'Summarize', 'Translate', 'Visualize', 'Acquire', 'Adapt', 'Allocate', 'Alphabetize', 'Apply', 'Ascertain', 'Assign', 'Attain', 'Avoid', 'Back up', 'Calculate', 'Capture', 'Change', 'Complete', 'Construct', 
'Customize', 'Demonstrate', 'Depreciate', 'Derive', 'Determine', 'Diminish', 'Discover', 'Employ', 'Examine', 'Exercise', 'Explore', 'Expose', 'Figure', 'Graph', 'Handle', 'Illustrate', 'Interconvert', 'Investigate', 'Manipulate', 'Modify', 
'Operate', 'Personalize', 'Plot','Practice', 'Prepare', 'Price', 'Process', 'Produce', 'Project', 'Provide', 'Relate', 'Round off', 'Sequence', 'Show', 'Simulate', 'Sketch', 'Solve', 'Subscribe', 'Transcribe', 'Use', 'Analyze', 'Audit', 
'Blueprint', 'Breadboard', 'Break down', 'Confirm', 'Correlate', 'Detect', 'Diagnose', 'Diagram', 'Discriminate', 'Dissect', 'Document', 'Ensure', 'Figure out', 'File', 'Group', 'Interrupt', 'Inventory', 'Layout', 'Manage', 'Maximize', 
'Minimize', 'Optimize', 'Order', 'Point out', 'Prioritize', 'Proofread', 'Query', 'Separate', 'Subdivide', 'Train', 'Transform', 'Appraise', 'Assess', 'Conclude', 'Counsel', 'Criticize', 'Critique', 'Evaluate', 'Grade', 'Hire', 'Judge', 
'Justify', 'Measure', 'Prescribe', 'Rank', 'Rate', 'Recommend', 'Release', 'Support', 'Test', 'Validate', 'Verify', 'Abstract', 'Animate', 'Arrange', 'Assemble', 'Budget', 'Categorize', 'Code', 'Combine', 'Compile', 'Compose', 'Cope', 
'Correspond', 'Create', 'Cultivate', 'Debug', 'Depict', 'Design', 'Develop', 'Devise', 'Dictate', 'Enhance', 'Facilitate', 'Format', 'Formulate', 'Generate', 'Import', 'Improve', 'Incorporate', 'Integrate', 'Interface', 'Join', 'Lecture', 
'Model', 'Network', 'Organize', 'Overhaul', 'Plan', 'Portray', 'Program', 'Rearrange', 'Reconstruct', 'Reorganize', 'Revise', 'Specify']))
    # Concat the generalised list of words to the PLO list
    train_column  = np.concatenate((feature_names_plo, append_words))
    # CLO list of words
    test_column = feature_names_clo
    # Intersection method for extracting common column names from the tables (both CLO AND PLO)
    # This is the column names from both the tables (using intersection)
    common_column = list(set(train_column).intersection(set(test_column)))
    # Filter the common column values from the CLO table
    U3_common = U2_clo.filter(list(common_column), axis=1)
    Cs = []
    Dds = []
    Ds = []
    # Extracting first row from PLO table and make a dataframe
    for x in range(len(df_plo)):
        Cs.append(U1_plo.loc[['P'+str(x+1)]])
    # Concatenating these extracted each PLOs with 'n' number of CLOs
    for x in range(len(df_plo)):
        Dds.append(pd.concat([Cs[x],U3_common], sort=True))
    # Filling the non values of the concatenated dataframes
    for x in range(len(df_plo)):
        Ds.append(Dds[x].fillna(0))
    # Calculate cosine similarity for concatenated dataframes and create a new dataframe
    for x in range(len(df_plo)):
        Dds[x] = pd.DataFrame(cosine_similarity(Ds[x], dense_output=True))
    # Extract the '0'th column because it has the CLO-PLO  cosine similarity values. We are neglecting the remaining ones.
    # Renaming the '0'th column name to 'Pn' ['P1, P2, P3, P4, ... 'Pn']
    for x in range(len(df_plo)):
        Dds[x].rename(columns = {0 :'P'+str(x+1)}, inplace = True)
    # Concatenating each  '0'th column from different cosine similarity dataframes
    Ddn = []
    for x in range(len(df_plo)):
        Ddn.append(Dds[x]['P'+str(x+1)])
    dd = pd.concat(Ddn, axis=1)[1:]
    # resetting index
    dd.reset_index(inplace = True)
    dd.drop(['index'], axis=1, inplace = True)
    # Setting threshold value 
    # Taking min max average of each column and set that as a threshold value
    for x in range(len(df_plo)):
        tes = dd['P'+str(x+1)].values.min()
        tes1 = dd['P'+str(x+1)].values.max()
        tt1 = (tes+tes1)/2
        if tt1 == 0:
            dd['P'+str(x+1)] = dd['P'+str(x+1)] 
        else:
            dd['P'+str(x+1)] = dd['P'+str(x+1)].apply(lambda x: 1 if x >= tt1 else 0)
    df_new = pd.concat([df_clo[1], dd], axis=1).set_index(1)
    df_new.to_csv('PLO-CLO.csv', index = True)
    global global_df
    global_df = df_new.copy()
    print(df_new)
    return render_template('table.html', df=df_new)

@app.route("/download", methods=['POST'])
def download():
    # Generate CSV data from Pandas DataFrame
    # pd.read_csv(temp_file.name, index=True)
    csv_file = global_df.to_csv(index=True).encode()
    # Convert CSV data to file object
    csv_file = io.BytesIO(csv_file)
    # Return CSV file as a download
    return send_file(csv_file, as_attachment=True, download_name='data.csv', mimetype='text/csv')

@app.route("/plo_clo_yes", methods=['GET', 'POST'])
def course_mapping():
    if request.method == 'POST':
        file = request.files['file']
        if file and file.filename.endswith('.docx'):
            # filename = secure_filename(file.filename)
            file_bytes = file.read()
            document_clo = Document(io.BytesIO(file_bytes)) # Creating a dataframe for CLOs and it will accept 'n' number of CLOs
            print(type(document_clo))
            
            # Connect to the SQLite database
            conn = sqlite3.connect('db/database.db')
            # Retrieve the blob data from the database
            cursor = conn.cursor()
            cursor.execute("SELECT program_id FROM Course WHERE course_name=?", (session['course_name'],))
            prog_id = cursor.fetchone()[0]
            cursor.execute("SELECT plo_clo FROM Program WHERE id=?", (prog_id,))
            blob_data = cursor.fetchone()[0]
            bytes_io = io.BytesIO(blob_data)
            document = Document(bytes_io) # extract tables from word document
            table_num=1
            df_plo = read_docx_table(document,table_num) # Creating a dataframe for PLOs and it will accept 'n' number of PLOs
            plo_copy = df_plo.copy()
            count_vectorizer = CountVectorizer(stop_words='english', min_df=0.005) # assigning count vectorizer breaking down a text into words
            # Data preprocessing for PLO dataframe
            plo_copy[1] = plo_copy[1].str.lower()
            corpus_plo = plo_copy[1].tolist() #it will seperate the rows into pairs 
            corpii_plo = count_vectorizer.fit_transform(corpus_plo)
            # extracting features names from PLO table
            feature_names_plo = count_vectorizer.get_feature_names_out()
            # Converting features to vector form and create a dataframe
            X1_plo = pd.DataFrame(corpii_plo.toarray(), columns=feature_names_plo)
            
            # Creating a dataframe for CLOs and it will accept 'n' number of CLOs
            table_num = 1
            # Close the database connection
            conn.close()
            df_clo = read_docx_table(document_clo,table_num)
            clo_copy = df_clo.copy()
            # Data preprocessing for CLO dataframe
            clo_copy[1] = clo_copy[1].str.lower()
            corpus_clo = clo_copy[1].tolist()
            corpii_clo = count_vectorizer.fit_transform(corpus_clo)
            # extracting features names from CLO table
            feature_names_clo = count_vectorizer.get_feature_names_out()
            # Converting features to vector form and create a dataframe
            X2_clo = pd.DataFrame(corpii_clo.toarray(), columns=feature_names_clo)
            # adding column index to the PLO table
            U1_plo = pd.concat([df_plo[0], X1_plo], axis=1).set_index(0)
            # adding column index to the CLO table
            U2_clo = pd.concat([df_clo[0], X2_clo], axis=1).set_index(0)
            append_words = list(map(str.lower,['Cite', 'Define', 'Describe', 'Draw', 'Enumerate', 'Identify' 'Index', 'Indicate', 'Label', 'List', 'Match', 'Meet', 'Name', 'Outline', 'Point', 'Quote', 'Read', 'Recall', 'Recite', 'Recognize', 'Record', 'Repeat', 'Reproduce','Review',
        'Select', 'State', 'Study', 'Tabulate', 'Trace', 'Write', 'Add', 'Approximate', 'Articulate', 'Associate', 'Characterize', 'Clarify', 'Classify', 'Compare', 'Compute', 'Contrast', 'Convert', 'Defend', 'Detail', 'Differentiate',
        'Discuss', 'Distinguish', 'Elaborate', 'Estimate', 'Example', 'Explain', 'Express', 'Extend', 'Extrapolate', 'Factor', 'Generalize', 'Give', 'Infer', 'Interact', 'Interpolate', 'Interpret', 'Observe', 'Paraphrase', 'Picture graphically',
        'Predict', 'Rewrite', 'Subtract', 'Summarize', 'Translate', 'Visualize', 'Acquire', 'Adapt', 'Allocate', 'Alphabetize', 'Apply', 'Ascertain', 'Assign', 'Attain', 'Avoid', 'Back up', 'Calculate', 'Capture', 'Change', 'Complete', 'Construct', 
        'Customize', 'Demonstrate', 'Depreciate', 'Derive', 'Determine', 'Diminish', 'Discover', 'Employ', 'Examine', 'Exercise', 'Explore', 'Expose', 'Figure', 'Graph', 'Handle', 'Illustrate', 'Interconvert', 'Investigate', 'Manipulate', 'Modify', 
        'Operate', 'Personalize', 'Plot','Practice', 'Prepare', 'Price', 'Process', 'Produce', 'Project', 'Provide', 'Relate', 'Round off', 'Sequence', 'Show', 'Simulate', 'Sketch', 'Solve', 'Subscribe', 'Transcribe', 'Use', 'Analyze', 'Audit', 
        'Blueprint', 'Breadboard', 'Break down', 'Confirm', 'Correlate', 'Detect', 'Diagnose', 'Diagram', 'Discriminate', 'Dissect', 'Document', 'Ensure', 'Figure out', 'File', 'Group', 'Interrupt', 'Inventory', 'Layout', 'Manage', 'Maximize', 
        'Minimize', 'Optimize', 'Order', 'Point out', 'Prioritize', 'Proofread', 'Query', 'Separate', 'Subdivide', 'Train', 'Transform', 'Appraise', 'Assess', 'Conclude', 'Counsel', 'Criticize', 'Critique', 'Evaluate', 'Grade', 'Hire', 'Judge', 
        'Justify', 'Measure', 'Prescribe', 'Rank', 'Rate', 'Recommend', 'Release', 'Support', 'Test', 'Validate', 'Verify', 'Abstract', 'Animate', 'Arrange', 'Assemble', 'Budget', 'Categorize', 'Code', 'Combine', 'Compile', 'Compose', 'Cope', 
        'Correspond', 'Create', 'Cultivate', 'Debug', 'Depict', 'Design', 'Develop', 'Devise', 'Dictate', 'Enhance', 'Facilitate', 'Format', 'Formulate', 'Generate', 'Import', 'Improve', 'Incorporate', 'Integrate', 'Interface', 'Join', 'Lecture', 
        'Model', 'Network', 'Organize', 'Overhaul', 'Plan', 'Portray', 'Program', 'Rearrange', 'Reconstruct', 'Reorganize', 'Revise', 'Specify']))
            # Concat the generalised list of words to the PLO list
            train_column  = np.concatenate((feature_names_plo, append_words))
            # CLO list of words
            test_column = feature_names_clo
            # Intersection method for extracting common column names from the tables (both CLO AND PLO)
            # This is the column names from both the tables (using intersection)
            common_column = list(set(train_column).intersection(set(test_column)))
            # Filter the common column values from the CLO table
            U3_common = U2_clo.filter(list(common_column), axis=1)
            Cs = []
            Dds = []
            Ds = []
            # Extracting first row from PLO table and make a dataframe
            for x in range(len(df_plo)):
                Cs.append(U1_plo.loc[['P'+str(x+1)]])
            # Concatenating these extracted each PLOs with 'n' number of CLOs
            for x in range(len(df_plo)):
                Dds.append(pd.concat([Cs[x],U3_common], sort=True))
            # Filling the non values of the concatenated dataframes
            for x in range(len(df_plo)):
                Ds.append(Dds[x].fillna(0))
            # Calculate cosine similarity for concatenated dataframes and create a new dataframe
            for x in range(len(df_plo)):
                Dds[x] = pd.DataFrame(cosine_similarity(Ds[x], dense_output=True))
            # Extract the '0'th column because it has the CLO-PLO  cosine similarity values. We are neglecting the remaining ones.
            # Renaming the '0'th column name to 'Pn' ['P1, P2, P3, P4, ... 'Pn']
            for x in range(len(df_plo)):
                Dds[x].rename(columns = {0 :'P'+str(x+1)}, inplace = True)
            # Concatenating each  '0'th column from different cosine similarity dataframes
            Ddn = []
            for x in range(len(df_plo)):
                Ddn.append(Dds[x]['P'+str(x+1)])
            dd = pd.concat(Ddn, axis=1)[1:]
            # resetting index
            dd.reset_index(inplace = True)
            dd.drop(['index'], axis=1, inplace = True)
            # Setting threshold value 
            # Taking min max average of each column and set that as a threshold value
            for x in range(len(df_plo)):
                tes = dd['P'+str(x+1)].values.min()
                tes1 = dd['P'+str(x+1)].values.max()
                tt1 = (tes+tes1)/2
                if tt1 == 0:
                    dd['P'+str(x+1)] = dd['P'+str(x+1)] 
                else:
                    dd['P'+str(x+1)] = dd['P'+str(x+1)].apply(lambda x: 1 if x >= tt1 else 0)
            df_new = pd.concat([df_clo[1], dd], axis=1).set_index(1)
            df_new.to_csv('PLO-CLO.csv', index=True)
            print(df_new)
            global global_df
            global_df = df_new.copy()
            # Process the document as needed
            return render_template('table.html', df=df_new)
        else:
            return 'Only .docx files are allowed'
    else:
        return render_template('upload.html')



####################################################################################################################
@app.route('/CLOandPLO', methods=['GET', 'POST']) 
@is_faculty_logged_in 
def CLOandPLO(): 
    std_reg = False 
    std = Student.query.all() 
    try: 
        if session['is_admin']: 
            courses = Course.query.all() 
        else: 
            courses = Course.query.filter_by(teacher_email=session['email']).all() 
    except: 
        courses = Course.query.filter(Course.students_email.contains(session['email'])).all() 
 
    if len(std) > 0: 
        std_reg = True 

    if request.method == 'POST': 
        
        session['course_id'] = request.form['course'] 
        course = Course.query.filter_by(course_id=request.form['course']).one() 
        session['course_name'] = course.course_name 
        

        # Handle the user's choice
        mapping_choice = request.form.get("clo")
       
        if mapping_choice == 'yes': 
           
            return redirect(url_for('course_mapping'))

        elif mapping_choice == 'no':
            return redirect(url_for('program_mapping'))

    return render_template('CLOandPLO.html', std_reg=std_reg, courses=courses)
##################################################################################################
#ATTENTION
##################################################################################################

@app.route('/attention', methods=['GET', 'POST'])
def mark_attention_1():
    std_reg = False
    std = Student.query.all()
    try:
        if session['is_admin']:
            courses = Course.query.all()
        else:
            courses = Course.query.filter_by(teacher_email=session['email']).all()
    except:
        courses = Course.query.filter(Course.students_email.contains(session['email'])).all()

    if len(std) > 0:
        std_reg = True
    if request.method == 'POST':
        session['lecture_no'] = request.form['lecture']
        session['course_id'] = request.form['course']
        course = Course.query.filter_by(course_id=request.form['course']).one()
        session['course_name'] = course.course_name
        return redirect(url_for('mark_attention_2'))

    return render_template('mark_attention_1.html', std_reg=std_reg, courses=courses)


@app.route('/attention/fr')
def mark_attention_2():
    # courses = Course.query.filter_by(teacher_email=session['email']).all()
    
    return render_template('mark_attention_2.html')


@app.route('/view_lectures_attention/', methods=['GET', 'POST'])
@is_faculty_logged_in
def view_lectures_attendance():
    if session['is_admin']:
        courses = Course.query.all()
    else:
        courses = Course.query.filter_by(teacher_email=session['email']).all()
    print(courses)

    if request.method=="POST":
        course_id = request.form['course']
        lect_no = request.form['lecture']
        if lect_no is None:
            lect_no = 1
        
        print(lect_no, course_id, session['name'])
        # get the attendance for the specified lecture
        attendances = Attendance.query.filter(Attendance.lecture_no==lect_no, Attendance.course_id==course_id).all()
        
        attendance_details = []
        for attendance in attendances:
            student = Student.query.filter_by(student_id=attendance.student_id).one()
            attendance_details.append(
                {
                    "student_id": attendance.student_id,
                    "student_name": student.name,
                    "lecture_no" : attendance.lecture_no,
                    "marked_by" : attendance.marked_by,
                    "marked_date" : attendance.marked_date,
                    "marked_time" : attendance.marked_time
                }
            )                  
                    
        selected_course = Course.query.filter_by(course_id=course_id).one()
        course_name = selected_course.course_name

        return render_template('view_lecture_attention.html', lect_no=lect_no, attendance=attendance_details, courses=courses, course_name=course_name)

    
    return render_template('view_lecture_attention.html', courses=courses)

# TODO download attendance for each lecture as marked by the teacher

@app.route('/facultydownloadcsv', defaults={'lect_no': 1})
@app.route('/facultydownloadcsv/<int:lect_no>/')
@is_faculty_logged_in
def download_attendance_csv(lect_no):
    headings = 'student_id,Course,Lecture_no,Marked_by,Marking_date,Marking_Time\n'
    attendance = Attendance.query.filter_by(lecture_no=lect_no).all()
    rows = ''
    for a in attendance:
        course = Course.query.filter_by(course_id=a.course_id).first()

        rows += str(a.student_id)+','+str(course.course_name)+','+str(a.lecture_no)+',' + \
            (a.marked_by)+','+str(a.marked_date)+','+str(a.marked_time)+' \n'
    csv = headings+rows
    return Response(
        csv,
        mimetype="text/csv",
        headers={"Content-disposition":
                "attachment; filename=attendance.csv"})


@app.route('/fr_attention')
def mark_face_attention():
    video_capture = VideoCapture(0)
    known_face_encodings = []
    known_face_names = []
    known_faces_filenames = []
    for (dirpath, dirnames, filenames) in os.walk('static/images/users'):
        known_faces_filenames.extend(filenames)
        break
    for filename in known_faces_filenames:
        face = face_recognition.load_image_file(
            'static/images/users/' + filename)
        known_face_names.append(filename[:-4])
        known_face_encodings.append(face_recognition.face_encodings(face)[0])
    face_locations = []
    face_encodings = []
    face_names = []
    process_this_frame = True
    total_item = []
    while True:
        frame = video_capture.read()
        date=datetime.date(datetime.now()),
        time=datetime.time(datetime.now())
        # Process every frame only one time
        if process_this_frame:
            # Find all the faces and face encodings in the current frame of video
            face_locations = face_recognition.face_locations(frame)
            face_encodings = face_recognition.face_encodings(
                frame, face_locations)
            # Initialize an array for the name of the detected users
            face_names = []
            # * ---------- Initialyse JSON to EXPORT --------- *
            json_to_export = {}
            flag = False
            for face_encoding in face_encodings:
                # See if the face is a match for the known face(s)
                matches = face_recognition.compare_faces(
                    known_face_encodings, face_encoding)
                roll_name = "Unknown"  # roll_name variables has roll no. and named saved. e.g. rahul-1666
                # Use the known face with the smallest distance to the new face
                face_distances = face_recognition.face_distance(
                    known_face_encodings, face_encoding)
                best_match_index = np.argmin(face_distances)
                if matches[best_match_index]:
                    roll_name = known_face_names[best_match_index]
                    roll = roll_name.split('-')[0]
                    # Check if the entry is is already in the DB
                    exists = Attendance.query.filter_by(student_id=roll).first()
                    # Check if the student's attendance is already marked
                    print(exists, type(exists))
                    if exists:
                        is_marked_already = Attendance.query.filter(
                            Attendance.student_id == roll, Attendance.course_id == session['course_id'], Attendance.lecture_no == int(session['lecture_no'])).all()
                    # If this student is not in attendance, create an entry
                    if exists is None:
                        # Create a new row for the student:
                        attendance = Attendance(
                            student_id=roll,
                            course_id=session['course_id'],
                            lecture_no=session['lecture_no'],
                            marked_by=session['name'],
                            marked_date=datetime.date(datetime.now()),
                            marked_time=datetime.time(datetime.now())
                        )
                        db.session.add(attendance)
                        db.session.commit()
                        flag = True
                        print("new entry")
                    # else if the student is already in the attendance, then check if his attendance is already marked for current lecture
                    elif len(is_marked_already) == 0:
                        # Create a new row for the student:
                        attendance = Attendance(
                            student_id=roll,
                            course_id=session['course_id'],
                            lecture_no=session['lecture_no'],
                            marked_by=session['name'],
                            marked_date=datetime.date(datetime.now()),
                            marked_time=datetime.time(datetime.now())
                        )
                        db.session.add(attendance)
                        db.session.commit()
                        flag = True
                        print("existing entry")
                        
                face_names.append(roll_name)
        process_this_frame = not process_this_frame
        # Display the results
        for (top, right, bottom, left), roll_name in zip(face_locations, face_names):
            # Draw a box around the face
            cv2.rectangle(frame, (left, top), (right, bottom), (0, 0, 255), 2)
            # Draw a label with a name below the face
            font = cv2.FONT_HERSHEY_DUPLEX
            cv2.putText(frame, roll_name, (left + 6, bottom - 6),
                        font, 0.5, (255, 255, 255), 1)
            if flag:
                cv2.putText(frame, 'Marked', (left + 12, bottom - 12),
                            font, 0.5, (255, 255, 255), 1)
        # Preprocess image for detection
        img = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        img = cv2.resize(img, (img_size, img_size), interpolation=cv2.INTER_LINEAR)
        results = model(img)
        detect_res = results.pandas().xyxy[0].to_json(orient="records")  # JSON img1 predictions
        detect_res = json.loads(detect_res)
        for res in detect_res:
            item = []
            if res['confidence']>conf_thresh:
                item.extend([date, time, res['name'], res['confidence']])
                total_item.append(item)
        # Display the resulting image
        cv2.imshow('Marking attendance', frame)
        # Hit 'q' on the keyboard to quit!
        if cv2.waitKey(1) & 0xFF == ord('q'):
            break
    report = pd.DataFrame(total_item, columns = ['Date', 'Time', 'Name', 'Confidence'])
    report.to_csv('results.csv', encoding='utf-8')
    # Release handle to the webcam
    video_capture.release()
    cv2.destroyAllWindows()
    return redirect(url_for('mark_attention_2'))

if __name__ == '__main__':
    
    app.run(debug=True)
